"""
Script de corrección de informes de auditoría.
Convierte los 5 DOCX de borradores de trabajo a informes profesionales para entrega a cliente.
"""

import sys
import re
import copy
from docx import Document
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

# ─────────────────────────────────────────────
# Utilidades generales
# ─────────────────────────────────────────────

def para_text(para):
    return ''.join(r.text for r in para.runs)

def cell_text(cell):
    return '\n'.join(para_text(p) for p in cell.paragraphs)

def replace_in_para(para, old, new):
    """Reemplaza texto en un párrafo preservando la estructura de runs."""
    full = para_text(para)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    # Colocar todo en el primer run, limpiar el resto
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ''
    else:
        # Párrafo sin runs — modificar el XML directamente
        para.text = new_full
    return True

def replace_in_cell(cell, old, new):
    changed = False
    for para in cell.paragraphs:
        if replace_in_para(para, old, new):
            changed = True
    return changed

def replace_everywhere(doc, replacements):
    """Aplica lista de (old, new) a todo el documento: párrafos + tablas."""
    for para in doc.paragraphs:
        for old, new in replacements:
            replace_in_para(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old, new in replacements:
                    replace_in_cell(cell, old, new)

def delete_para(para):
    p = para._element
    p.getparent().remove(p)

def delete_table_row(row):
    tr = row._element
    tr.getparent().remove(tr)

def remove_section(doc, start_marker, end_marker, inclusive_end=False):
    """
    Elimina todos los párrafos desde el que contiene start_marker
    hasta (sin incluir) el que contiene end_marker.
    Si inclusive_end=True también elimina el párrafo con end_marker.
    """
    paras = list(doc.paragraphs)
    to_delete = []
    in_section = False
    for para in paras:
        txt = para_text(para)
        if start_marker in txt and not in_section:
            in_section = True
        if in_section:
            if end_marker in txt and txt.strip() != '' and not (start_marker in txt):
                if inclusive_end:
                    to_delete.append(para)
                break
            to_delete.append(para)
    for p in to_delete:
        delete_para(p)
    return len(to_delete)

def strip_emoji_prefix(para, emoji):
    """Elimina un emoji del inicio del texto de un párrafo."""
    txt = para_text(para)
    if txt.strip().startswith(emoji):
        replace_in_para(para, emoji, '')

def find_table_row_containing(doc, text):
    """Busca la fila de tabla que contiene el texto dado."""
    for table in doc.tables:
        for row in table.rows:
            row_text = ' '.join(cell_text(c) for c in row.cells)
            if text in row_text:
                return row
    return None

def replace_in_table_cell_matching(doc, row_text_marker, old, new):
    """En la fila que contiene row_text_marker, reemplaza old por new."""
    for table in doc.tables:
        for row in table.rows:
            rt = ' '.join(cell_text(c) for c in row.cells)
            if row_text_marker in rt:
                for cell in row.cells:
                    replace_in_cell(cell, old, new)

# ─────────────────────────────────────────────
# Reemplazos comunes a TODOS los documentos
# ─────────────────────────────────────────────

COMMON_REPLACEMENTS = [
    # Quitar marcadores de revisión interna del encabezado
    ('— REVISADO POST-VALIDACIÓN FORENSE', ''),
    ('— REVISADO POST-VALIDACIÓN FORENSE 11/05/2026', ''),
    ('REVISADO POST-VALIDACIÓN FORENSE 11/05/2026', ''),
    ('REVISADO POST-VALIDACIÓN FORENSE', ''),
    # Quitar timestamps de trabajo
    ('(rev. 11/05/2026 16:50)', ''),
    ('(rev. 16:50)', ''),
    ('rev. 11/05/2026 16:50', ''),
    # Quitar marcadores de proceso interno
    (' ⭐ NUEVO', ''),
    ('⭐ NUEVO', ''),
    (' ⭐', ''),
    ('⭐', ''),
    # Quitar 🆕 como marcador de proceso (se reemplaza abajo en contexto específico)
    # IP del servidor (información sensible de infraestructura)
    ('192.168.1.51:1433', 'servidor ERP S10 productivo'),
    ('192.168.1.51 , base CMO', 'sistema ERP S10 productivo'),
    ('192.168.1.51', 'sistema ERP S10 productivo'),
    (' , base de datos CMO con 2,035 tablas + 1,382 vistas', ''),
]

# ─────────────────────────────────────────────
# DOC 01 — RESUMEN_GERENCIAL
# ─────────────────────────────────────────────

def corregir_doc01():
    doc = Document('01_RESUMEN_GERENCIAL.docx')

    # 1. Reemplazos comunes
    replace_everywhere(doc, COMMON_REPLACEMENTS)

    # 2. Reemplazos específicos Doc01
    especificos = [
        # Lenguaje informal
        ('sin excusas', 'con carácter urgente'),
        # Quitar referencia al doc interno de S10 en contexto ejecutivo
        ('Ver 05_VALIDACION_FORENSE_S10.docx para el detalle metodológico.', ''),
        # Limpiar marcadores 🆕 de hallazgos nuevos en intro
        ('🆕 7 hallazgos nuevos detectados por la validación: ver §2.4.', ''),
        ('🆕 7 hallazgos nuevos detectados por la validación: ver §2.4', ''),
        # Quitar "⭐ NUEVO" en referencias a anexo
        ('Anexo 4: Validación forense contra S10 productivo — finding-by-finding ⭐ NUEVO',
         'Anexo 4: Validación forense contra S10 productivo — finding-by-finding'),
        # Nota técnica interna de proceso
        ('Ver 05_VALIDACION_FORENSE_S10.docx para el detalle metodológico', ''),
    ]
    replace_everywhere(doc, especificos)

    # 3. Eliminar Sección 0 "NOTA DE REVISIÓN" completa
    #    Empieza en "0. NOTA DE REVISIÓN", termina antes de "1. RESUMEN EJECUTIVO"
    eliminados = remove_section(doc, '0. NOTA DE REVISIÓN', '1. RESUMEN EJECUTIVO', inclusive_end=False)
    print(f'  Doc01: eliminados {eliminados} párrafos de Sección 0')

    # 4. Limpiar cualquier párrafo que quedó vacío por las eliminaciones anteriores
    #    (solo los que estén cerca de donde estaba la sección 0)
    for para in doc.paragraphs:
        if para_text(para).strip() in ('', ' '):
            # Dejar algunos para separación; solo eliminar si son consecutivos al área
            pass

    doc.save('01_RESUMEN_GERENCIAL.docx')
    print('  ✓ Doc01 guardado')

# ─────────────────────────────────────────────
# DOC 02 — ANEXOS_REGISTROS
# ─────────────────────────────────────────────

def corregir_doc02():
    doc = Document('02_ANEXOS_REGISTROS.docx')

    replace_everywhere(doc, COMMON_REPLACEMENTS)

    especificos = [
        # Quitar ⭐ del índice de anexos
        (' ⭐ CMO', ' CMO'),
        ('Q ⭐', 'Q'),
        ('R ⭐', 'R'),
        ('S ⭐', 'S'),
        # Aclarar Anexo H
        ('Descuadre operativo por NroD', 'Diferencia Db−Cr real por NroD (descuadre contable, no suma de líneas)'),
    ]
    replace_everywhere(doc, especificos)

    # Limpiar cualquier 🆕 residual
    for para in doc.paragraphs:
        replace_in_para(para, '🆕', 'NUEVO:')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para, '🆕', 'NUEVO:')

    doc.save('02_ANEXOS_REGISTROS.docx')
    print('  ✓ Doc02 guardado')

# ─────────────────────────────────────────────
# DOC 03 — INFORME_CONSOLIDADO_FINAL
# ─────────────────────────────────────────────

def corregir_doc03():
    doc = Document('03_INFORME_CONSOLIDADO_FINAL.docx')

    replace_everywhere(doc, COMMON_REPLACEMENTS)

    # 1. Eliminar Sección 0 "NOTA DE REVISIÓN POST-VALIDACIÓN"
    eliminados = remove_section(doc,
        'NOTA DE REVISIÓN POST-VALIDACIÓN',
        '1. ALCANCE',
        inclusive_end=False)
    if eliminados == 0:
        # Intentar con variante sin número
        eliminados = remove_section(doc,
            '0. NOTA',
            '1.',
            inclusive_end=False)
    print(f'  Doc03: eliminados {eliminados} párrafos de Sección 0')

    # 2. Eliminar fila CMO-05 "REFUTADO" de la tabla de hallazgos
    row_cmo05 = find_table_row_containing(doc, 'CMO-05')
    if row_cmo05:
        delete_table_row(row_cmo05)
        print('  Doc03: eliminada fila CMO-05 (hallazgo refutado)')
    else:
        print('  Doc03: fila CMO-05 no encontrada (puede estar en otro contexto)')

    # 3. MEDARQ-01: cambiar severidad CRÍTICO → OBSERVACIÓN y 🔴 → 🟡
    #    en la fila que contiene "MEDARQ-01"
    for table in doc.tables:
        for row in table.rows:
            row_txt = ' '.join(cell_text(c) for c in row.cells)
            if 'MEDARQ-01' in row_txt and ('CRÍTICO' in row_txt or '🔴' in row_txt):
                for cell in row.cells:
                    replace_in_cell(cell, 'CRÍTICO', 'OBSERVACIÓN')
                    replace_in_cell(cell, '🔴', '🟡')
                print('  Doc03: corregida severidad MEDARQ-01 CRÍTICO→OBSERVACIÓN')
                break

    # 4. Corregir nota "IMPORTANTE" en tono defensivo
    replace_everywhere(doc, [
        ('IMPORTANTE: Las cifras', 'Nota aclaratoria: Las cifras'),
        ('"S/122.9M movido CMO"', 'S/122.9M movido CMO'),
        ('"S/123M"', 'S/123M'),
    ])

    # 5. Limpiar 🆕 residuales
    for para in doc.paragraphs:
        replace_in_para(para, '🆕', '')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para, '🆕', '')

    doc.save('03_INFORME_CONSOLIDADO_FINAL.docx')
    print('  ✓ Doc03 guardado')

# ─────────────────────────────────────────────
# DOC 04 — TRAZABILIDAD_HISTORICA
# ─────────────────────────────────────────────

def corregir_doc04():
    doc = Document('04_TRAZABILIDAD_HISTORICA.docx')

    replace_everywhere(doc, COMMON_REPLACEMENTS)

    especificos = [
        # Corrección factual crítica: BBVA "desde 2023" → "desde 2026"
        ('Confirmado real desde 2023 (241 movs)', 'Confirmado real desde 2026 (origen enero 2026)'),
        ('confirmado real desde 2023', 'confirmado real desde 2026'),
        ('arrancó en 2023', 'arrancó en 2026'),
        ('desde 2023 (241', 'desde 2026 (241'),
        # Actualizar base de contingencia en sección 7.2 (S/123M → S/11.1M)
        ('S/123M @ 29.5%', 'S/11.1M @ 29.5%'),
        ('S/122.9M', 'S/11.1M'),
        ('sobre S/123M', 'sobre S/11.1M'),
        ('base S/123M', 'base S/11.1M'),
        ('S/2,700,000', 'S/3,272,000'),
        # Primera persona → voz impersonal/plural
        ('Mi mayor error fue', 'El error de interpretación inicial fue'),
        ('Mi error inicial fue', 'El error inicial fue'),
        ('mi mayor error', 'el error de interpretación'),
        ('mi error', 'el error'),
        ('Mis conclusiones', 'Las conclusiones'),
        ('me equivoqué', 'la interpretación fue incorrecta'),
        # Quitar calificativo informal "errado" → más neutro
        ('Hallazgo inicial (errado):', 'Interpretación inicial (corregida):'),
        ('hallazgo inicial errado', 'interpretación inicial corregida'),
        # Limpiar 🆕 de secciones N01-N07
        ('🆕 N01', 'N01'),
        ('🆕 N02', 'N02'),
        ('🆕 N03', 'N03'),
        ('🆕 N04', 'N04'),
        ('🆕 N05', 'N05'),
        ('🆕 N06', 'N06'),
        ('🆕 N07', 'N07'),
        # Tono profesional
        ('el usuario pidió', 'se solicitó'),
        ('El usuario pidió', 'Se solicitó'),
    ]
    replace_everywhere(doc, especificos)

    doc.save('04_TRAZABILIDAD_HISTORICA.docx')
    print('  ✓ Doc04 guardado')

# ─────────────────────────────────────────────
# DOC 05 — VALIDACION_FORENSE_S10
# ─────────────────────────────────────────────

def corregir_doc05():
    doc = Document('05_VALIDACION_FORENSE_S10.docx')

    replace_everywhere(doc, COMMON_REPLACEMENTS)

    especificos = [
        # Quitar 🆕 de secciones de hallazgos nuevos
        ('🆕 N01 —', 'N01 —'),
        ('🆕 N02 —', 'N02 —'),
        ('🆕 N03 —', 'N03 —'),
        ('🆕 N04 —', 'N04 —'),
        ('🆕 N05 —', 'N05 —'),
        ('🆕 N06 —', 'N06 —'),
        ('🆕 N07 —', 'N07 —'),
        ('🆕 N01', 'N01'),
        ('🆕 N02', 'N02'),
        ('🆕 N03', 'N03'),
        ('🆕 N04', 'N04'),
        ('🆕 N05', 'N05'),
        ('🆕 N06', 'N06'),
        ('🆕 N07', 'N07'),
        # Formalizar lenguaje informal
        ('Esto NO es el dashboard. Es contra la fuente.',
         'Esta validación se ejecutó directamente contra el origen de datos S10, no contra el dashboard.'),
        ('el usuario pidió validar', 'se solicitó validar'),
        ('El usuario pidió', 'Se solicitó'),
        # Validaciones con iconos de veredicto — mantener como están (🔴 🟡 🟢 son de contenido)
        # Solo limpiar las referencias de proceso
        ('s10-agent/validation-agent.js', 'el agente de validación forense'),
        ('s10-agent/validation-output.json (276 KB)', 'el archivo de salida de validación'),
        ('276 KB', ''),
    ]
    replace_everywhere(doc, especificos)

    doc.save('05_VALIDACION_FORENSE_S10.docx')
    print('  ✓ Doc05 guardado')

# ─────────────────────────────────────────────
# Ejecución
# ─────────────────────────────────────────────

if __name__ == '__main__':
    print('Iniciando corrección de informes de auditoría...\n')

    try:
        print('DOC 01 — Resumen Gerencial')
        corregir_doc01()
    except Exception as e:
        print(f'  ERROR Doc01: {e}')

    try:
        print('DOC 02 — Anexos y Registros')
        corregir_doc02()
    except Exception as e:
        print(f'  ERROR Doc02: {e}')

    try:
        print('DOC 03 — Informe Consolidado Final')
        corregir_doc03()
    except Exception as e:
        print(f'  ERROR Doc03: {e}')

    try:
        print('DOC 04 — Trazabilidad Histórica')
        corregir_doc04()
    except Exception as e:
        print(f'  ERROR Doc04: {e}')

    try:
        print('DOC 05 — Validación Forense S10')
        corregir_doc05()
    except Exception as e:
        print(f'  ERROR Doc05: {e}')

    print('\n✓ Proceso completado. Abre los documentos en Word para revisión final.')
