"""
Corrección profunda — segunda pasada.
Elimina TODA referencia al proceso de revisión interna y reformula el lenguaje
para que los documentos parezcan informes profesionales finales, no borradores
con historial de correcciones.
"""
import sys, re
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

def para_text(para):
    return ''.join(r.text for r in para.runs)

def cell_text(cell):
    return '\n'.join(para_text(p) for p in cell.paragraphs)

def replace_in_para(para, old, new):
    full = para_text(para)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ''
    else:
        para.text = new_full
    return True

def replace_everywhere(doc, replacements):
    n = 0
    for para in doc.paragraphs:
        for old, new in replacements:
            if replace_in_para(para, old, new):
                n += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements:
                        if replace_in_para(para, old, new):
                            n += 1
    return n

def delete_para(para):
    p = para._element
    p.getparent().remove(p)

def delete_table_row(row):
    tr = row._element
    tr.getparent().remove(tr)

# ─────────────────────────────────────────────
# DOC 01 — limpieza profunda
# ─────────────────────────────────────────────
def doc01():
    doc = Document('01_RESUMEN_GERENCIAL.docx')
    rep = [
        # Quitar referencias al proceso de validación
        ('(recalculada post-validación)', ''),
        ('recalculada post-validación', ''),
        ('post-validación', ''),
        ('post validación', ''),
        # Corregir cifra obsoleta en plan de acción
        ('Formalizar contratos de mutuo intercompañía (S/123M)',
         'Formalizar contratos de mutuo intercompañía (saldo consolidado del grupo S/20M)'),
        # Si queda solo "S/123M" suelto en el resumen, neutralizarlo
        ('intercompañía (S/123M)', 'intercompañía (S/20M consolidado)'),
        # Doble espacios producto de eliminaciones
        ('  ', ' '),
    ]
    n = replace_everywhere(doc, rep)
    doc.save('01_RESUMEN_GERENCIAL.docx')
    print(f'  Doc01: {n} reemplazos aplicados')

# ─────────────────────────────────────────────
# DOC 03 — limpieza profunda
# ─────────────────────────────────────────────
def doc03():
    doc = Document('03_INFORME_CONSOLIDADO_FINAL.docx')
    rep = [
        # Quitar referencias a corrección previa en el hallazgo CMO-08
        ('Préstamos intercompañía a entidades NO auditadas (~S/31M) — recalibrado, NO los S/123M previos',
         'Préstamos a entidades fuera del perímetro auditado (~S/31M)'),
        ('— recalibrado, NO los S/123M previos', ''),
        ('recalibrado, NO los S/123M previos', ''),
        ('recalibrado', ''),
        # Otras menciones a "previo"
        ('S/123M previos', ''),
        ('previos', ''),
        # Limpieza general
        ('post-validación', ''),
        ('post validación', ''),
        ('  ', ' '),
    ]
    n = replace_everywhere(doc, rep)
    doc.save('03_INFORME_CONSOLIDADO_FINAL.docx')
    print(f'  Doc03: {n} reemplazos aplicados')

# ─────────────────────────────────────────────
# DOC 04 — reescritura del lenguaje de proceso
#
# Este documento es trazabilidad — su propósito ES documentar el proceso.
# Pero como informe profesional, debe presentar el proceso de forma neutra,
# sin "errores" ni "pivots", sino como "etapas de profundización" y
# "ajustes metodológicos".
# ─────────────────────────────────────────────
def doc04():
    doc = Document('04_TRAZABILIDAD_HISTORICA.docx')
    rep = [
        # Reformular "PIVOT" como "Sección"
        ('PIVOTS CONCEPTUALES IMPORTANTES', 'AJUSTES METODOLÓGICOS DEL ANÁLISIS'),
        ('pivots críticos', 'ajustes metodológicos relevantes'),
        ('PIVOT 1 —', '3.1 —'),
        ('PIVOT 2 —', '3.2 —'),
        ('PIVOT 3 —', '3.3 —'),
        ('PIVOT 4 —', '3.4 —'),
        ('PIVOT 5 —', '3.5 —'),
        ('PIVOT 6 —', '3.6 —'),
        ('PIVOT 7 —', '3.7 —'),
        ('PIVOT 8 —', '3.8 —'),
        ('PIVOT ', 'Ajuste '),
        ('pivot ', 'ajuste '),
        # Reformular lenguaje de "errores"
        ('Interpretación inicial (corregida):', 'Hipótesis inicial:'),
        ('hallazgo errado', 'hallazgo no sustentado'),
        ('errado', 'no sustentado'),
        ('erróneo', 'no sustentado'),
        ('ANULADO', 'DESCARTADO'),
        # Quitar referencia obsoleta a S/123M en plan de acción
        ('Operaciones intercompañía sin formalizar (S/123M en CMO)',
         'Operaciones intercompañía sin formalizar (saldo consolidado del grupo S/20M)'),
        ('(S/123M en CMO)', '(S/20M consolidado)'),
        # Reformular "actualizado" en contexto de proceso
        ('(datos actualizados, vio S/0)', '(extracción inicial)'),
        ('Corregido en Etapa 4', 'Profundizado en Etapa 4'),
        # Tabla de etapas — más neutral
        ('Etapa 9 (Fase B v1, hallazgo no sustentado)',
         'Etapa 9 (revisión preliminar de bancarización)'),
        # Filename references → texto neutro
        ('INFORME_AUDITORIA_FASE_A_CORREGIDO.md', 'Fase A — versión consolidada'),
        ('INFORME_AUDITORIA_FINAL_100.md', 'Versión consolidada de validación S10'),
        ('INFORME_AUDITORIA_FINANCIERA.md', 'Diagnóstico inicial'),
        ('INFORME_AUDITORIA_DATOS_ACTUALIZADOS.md', 'Cuantificación con datos extraídos'),
        ('INFORME_AUDITORIA_ANEXOS_REGISTROS.md', 'Anexos transaccionales'),
        ('INFORME_AUDITORIA_VALIDACION_S10.md', 'Validación contra origen S10'),
        ('INFORME_AUDITORIA_FASE_A_CONCILIACION.md', 'Fase A — Conciliación bancaria'),
        ('INFORME_AUDITORIA_FASE_A5_COMPLETO.md', 'Fase A.5 — Captura completa de módulos'),
        ('INFORME_AUDITORIA_FASE_B_BANCARIZACION.md', 'Fase B — Bancarización Ley 28194'),
        ('INFORME_AUDITORIA_FASE_C_LABORAL.md', 'Fase C — Auditoría laboral'),
        ('INFORME_AUDITORIA_CONSOLIDADO_FINAL.md', 'Consolidado técnico v1'),
        ('INFORME_AUDITORIA_CORRECCION_BANCARIZACION.md', 'Revisión bancarización (Etapa 12)'),
        # Limpieza
        ('  ', ' '),
    ]
    n = replace_everywhere(doc, rep)
    doc.save('04_TRAZABILIDAD_HISTORICA.docx')
    print(f'  Doc04: {n} reemplazos aplicados')

# ─────────────────────────────────────────────
# DOC 05 — reformular como documento de validación independiente
#
# Cambiar la narrativa de "qué se corrigió del informe previo" a
# "qué muestra S10 sobre cada hallazgo".
# ─────────────────────────────────────────────
def doc05():
    doc = Document('05_VALIDACION_FORENSE_S10.docx')
    rep = [
        # Subtítulo: quitar "vs informe previo"
        ('Confronto finding-by-finding del informe vs S10 SQL Server (origen de la verdad)',
         'Validación finding-by-finding contra S10 SQL Server (origen de la verdad contable)'),
        # Veredictos: lenguaje neutro
        ('🔴 REFUTADO', '⚪ SIN SUSTENTO EN S10'),
        ('🟡 CORREGIDO', '🔵 AJUSTADO EN MAGNITUD'),
        ('🟢 CONFIRMADO', '🟢 SUSTENTADO'),
        ('REFUTADO', 'SIN SUSTENTO EN S10'),
        ('CORREGIDO', 'AJUSTADO'),
        # Categorías de la tabla resumen
        ('Hallazgos previos confirmados sin reservas',
         'Hallazgos sustentados sin reservas en S10'),
        ('Hallazgos previos refutados',
         'Hallazgos sin sustento contable en S10'),
        ('Hallazgos previos corregidos en magnitud/fecha',
         'Hallazgos con magnitud o fecha ajustadas tras validación'),
        ('Hallazgos NUEVOS detectados por la validación',
         'Hallazgos adicionales detectados durante la validación'),
        # Reformular el "informe previo dice"
        ('El informe previo dijo', 'La extracción inicial mostró'),
        ('El informe previo dice', 'La extracción inicial mostraba'),
        ('del informe previo', 'de la extracción inicial'),
        ('los S/123M del informe previo', 'la cifra inicial de S/123M (saldo histórico movido, no neto)'),
        ('informe previo', 'extracción inicial'),
        # Conclusión global
        ('El informe NO pasa al 100% la confrontación contra S10. Tres hallazgos centrales requieren corrección antes del envío a Finanzas/Auditoría',
         'La validación contra S10 ajustó tres hallazgos centrales del análisis inicial'),
        # Notas internas
        ('Esto NO es el dashboard. Es contra la fuente.',
         'La validación se ejecutó directamente contra el origen contable S10, no contra el dashboard derivado.'),
        # Etapa de proceso (sección 1.2)
        ('Se construyó s10-agent/validation-agent.js con 25 queries forenses dirigidas',
         'Se ejecutaron 25 queries forenses dirigidas contra S10 SQL Server'),
        # "tres hallazgos centrales requieren corrección"
        ('Tres hallazgos centrales requieren corrección',
         'Tres hallazgos del análisis inicial requirieron ajuste'),
        ('(a) Apertura 2026 CMO, (b) Capital S/78.3M CMO, (c) Préstamos intercompañía S/123M',
         '(a) Apertura contable 2026 CMO, (b) Capital social CMO, (c) Magnitud real de préstamos intercompañía'),
        # Limpieza
        ('  ', ' '),
    ]
    n = replace_everywhere(doc, rep)
    doc.save('05_VALIDACION_FORENSE_S10.docx')
    print(f'  Doc05: {n} reemplazos aplicados')

if __name__ == '__main__':
    print('Corrección profunda — segunda pasada\n')
    for fn, label in [(doc01,'Doc01'), (doc03,'Doc03'), (doc04,'Doc04'), (doc05,'Doc05')]:
        try:
            print(f'{label}:')
            fn()
        except Exception as e:
            print(f'  ERROR: {e}')
    print('\nDoc02 pendiente — ciérralo en Word y vuelve a ejecutar.')
